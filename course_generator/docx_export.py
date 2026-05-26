from pathlib import Path


def export_markdown_to_docx(markdown_text: str, output_path: Path) -> str:
    """Convert a simple Markdown summary into a DOCX file."""
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportError("python-docx is required for DOCX export. Run: pip install python-docx") from exc

    doc = Document()
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("**") and line.endswith("**"):
            doc.add_paragraph(line.strip("*").strip())
        else:
            doc.add_paragraph(line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return str(output_path)
