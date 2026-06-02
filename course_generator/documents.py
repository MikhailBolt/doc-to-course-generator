import hashlib
import json
from pathlib import Path
from typing import Any, Dict, List, NamedTuple, Optional

from course_generator.constants import SUPPORTED_EXTENSIONS
from course_generator.errors import DocumentSourceError
from course_generator.utils import clean_text


class DocCollection(NamedTuple):
    """Collected source files plus the root used for relative display names."""

    files: List[Path]
    root: Path
    truncated: bool = False
    total_found: int = 0


def document_display_name(file_path: Path, labels_base: Path) -> str:
    try:
        return str(file_path.resolve().relative_to(labels_base.resolve())).replace("\\", "/")
    except ValueError:
        return file_path.name


def collect_source_files(docs_path: str, *, recursive: bool = False, max_files: int | None = None) -> DocCollection:
    path = Path(docs_path)

    if not path.exists():
        raise DocumentSourceError(f"'{docs_path}' does not exist.")

    if path.is_file():
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise DocumentSourceError(f"'{docs_path}' is not a supported file type ({', '.join(sorted(SUPPORTED_EXTENSIONS))}).")
        return DocCollection(files=[path], root=path.parent.resolve(), truncated=False, total_found=1)

    root = path.resolve()

    def matching_files() -> List[Path]:
        if recursive:
            found: List[Path] = []
            for p in sorted(root.rglob("*")):
                if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS:
                    found.append(p)
            return found
        return sorted([p for p in root.iterdir() if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS])

    all_files = matching_files()
    if not all_files:
        hint = " (try --recursive-docs for subfolders)" if not recursive else ""
        raise DocumentSourceError(f"No supported files found in '{docs_path}'.{hint}")

    total_found = len(all_files)
    truncated = bool(max_files is not None and max_files > 0 and total_found > max_files)
    source_files = all_files[:max_files] if truncated else all_files

    return DocCollection(files=source_files, root=root, truncated=truncated, total_found=total_found)


def file_fingerprint(file_path: Path) -> str:
    stat = file_path.stat()
    raw = f"{file_path.resolve()}|{stat.st_size}|{stat.st_mtime}"
    return hashlib.sha256(raw.encode("utf-8")).hexdigest()


def build_manifest_data(source_files: List[Path]) -> Dict[str, Any]:
    return {
        "files": [
            {
                "name": f.name,
                "path": str(f.resolve()),
                "fingerprint": file_fingerprint(f),
            }
            for f in source_files
        ]
    }


def load_manifest(manifest_file: str) -> Dict[str, Any]:
    path = Path(manifest_file)
    if not path.exists():
        return {}
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_manifest(manifest_file: str, data: Dict[str, Any]) -> None:
    with open(manifest_file, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def is_index_stale(source_files: List[Path], db_path: str, manifest_file: str) -> bool:
    db_dir = Path(db_path)
    index_file = db_dir / "index.faiss"
    meta_file = db_dir / "index.pkl"

    if not db_dir.exists() or not index_file.exists() or not meta_file.exists():
        return True

    current_manifest = build_manifest_data(source_files)
    saved_manifest = load_manifest(manifest_file)
    return current_manifest != saved_manifest


def load_file_documents(file_path: Path, *, labels_base: Optional[Path] = None) -> List[Any]:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        from langchain_community.document_loaders import PyPDFLoader

        loader = PyPDFLoader(str(file_path))
        docs = loader.load()
    elif suffix in {".txt", ".md"}:
        from langchain_community.document_loaders import TextLoader

        loader = TextLoader(str(file_path), encoding="utf-8")
        docs = loader.load()
    elif suffix == ".docx":
        from langchain_core.documents import Document

        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise ImportError("python-docx is required for .docx input. Run: pip install python-docx") from exc

        docx = DocxDocument(str(file_path))
        paragraphs = [p.text.strip() for p in docx.paragraphs if p.text and p.text.strip()]
        text = "\n\n".join(paragraphs)
        if not text.strip():
            raise ValueError(f"DOCX file appears empty: {file_path.name}")
        docs = [Document(page_content=text, metadata={"source": str(file_path.resolve())})]
    else:
        raise ValueError(f"Unsupported file type: {suffix}")

    base = labels_base.resolve() if labels_base else None
    doc_name = document_display_name(file_path, base) if base else file_path.name

    for doc in docs:
        doc.metadata["document_name"] = doc_name
        doc.metadata["document_path"] = str(file_path.resolve())
        doc.metadata["document_type"] = suffix.lstrip(".")
        if base:
            doc.metadata["document_relative"] = doc_name
    return docs


def get_combined_preview_text(
    source_files: List[Path],
    *,
    labels_base: Optional[Path],
    max_chars_per_file: int = 6000,
) -> str:
    parts = []
    for file_path in source_files:
        label = document_display_name(file_path, labels_base) if labels_base else file_path.name
        try:
            docs = load_file_documents(file_path, labels_base=labels_base)
            joined = "\n".join(doc.page_content for doc in docs)
            joined = clean_text(joined)[:max_chars_per_file]
            parts.append(f"\n===== DOCUMENT: {label} =====\n{joined}\n")
        except Exception as exc:
            parts.append(f"\n===== DOCUMENT: {label} =====\nFailed to read document: {exc}\n")
    return "\n".join(parts)
